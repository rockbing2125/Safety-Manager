"""
Word文档解析器
"""
import sys
from pathlib import Path
from typing import List
from loguru import logger

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx未安装，Word解析功能不可用")

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))


class DocxParser:
    """Word文档解析器"""

    @staticmethod
    def is_available() -> bool:
        """检查Word解析功能是否可用"""
        return DOCX_SUPPORT

    @staticmethod
    def extract_text(docx_path: str) -> str:
        """提取Word文档文本"""
        if not DOCX_SUPPORT:
            return ""

        try:
            doc = Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text

        except Exception as e:
            logger.error(f"提取Word文本失败: {e}")
            return ""

    @staticmethod
    def extract_paragraphs(docx_path: str) -> List[str]:
        """提取Word文档段落"""
        if not DOCX_SUPPORT:
            return []

        try:
            doc = Document(docx_path)
            return [p.text for p in doc.paragraphs if p.text.strip()]

        except Exception as e:
            logger.error(f"提取Word段落失败: {e}")
            return []

    @staticmethod
    def get_paragraph_count(docx_path: str) -> int:
        """获取段落数量"""
        if not DOCX_SUPPORT:
            return 0

        try:
            doc = Document(docx_path)
            return len([p for p in doc.paragraphs if p.text.strip()])

        except Exception as e:
            logger.error(f"获取段落数量失败: {e}")
            return 0

    @staticmethod
    def get_core_properties(docx_path: str) -> dict:
        """获取Word文档核心属性"""
        if not DOCX_SUPPORT:
            return {}

        try:
            doc = Document(docx_path)
            props = doc.core_properties

            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }

        except Exception as e:
            logger.error(f"获取Word属性失败: {e}")
            return {}